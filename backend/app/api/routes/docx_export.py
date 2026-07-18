from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from fastapi import APIRouter, HTTPException
from fastapi.responses import Response

router = APIRouter()

# A4 (210mm) minus 15mm side margins — used to right-align dates via a tab stop.
USABLE_WIDTH_MM = 180

CENTER = WD_ALIGN_PARAGRAPH.CENTER
LEFT = WD_ALIGN_PARAGRAPH.LEFT

# Word equivalents of the on-screen templates: font family, header alignment,
# section-title treatment, accent colour, and spacing.
TEMPLATE_STYLES = {
    "original": {
        "font": "Arial", "body_size": 10.5, "name_size": 20, "title_size": 13, "contact_size": 9,
        "section_size": 11, "line_spacing": 1.15, "header_align": CENTER, "section_upper": False,
        "title_color": None, "section_color": None, "border_color": "000000", "border_sz": "6",
    },
    "modern": {
        "font": "Calibri", "body_size": 10.5, "name_size": 20, "title_size": 13, "contact_size": 9,
        "section_size": 11, "line_spacing": 1.15, "header_align": LEFT, "section_upper": True,
        "title_color": "2563EB", "section_color": "1F2937", "border_color": "2563EB", "border_sz": "12",
    },
    "classic": {
        "font": "Georgia", "body_size": 10.5, "name_size": 20, "title_size": 13, "contact_size": 9,
        "section_size": 11, "line_spacing": 1.25, "header_align": CENTER, "section_upper": False,
        "title_color": None, "section_color": None, "border_color": "000000", "border_sz": "6",
    },
    "compact": {
        "font": "Arial", "body_size": 9.5, "name_size": 18, "title_size": 12, "contact_size": 8.5,
        "section_size": 10, "line_spacing": 1.0, "header_align": LEFT, "section_upper": True,
        "title_color": None, "section_color": None, "border_color": "444444", "border_sz": "6",
    },
}


def _format_date_range(item: dict) -> str:
    sm = item.get("start_month") or ""
    sy = item.get("start_year") or ""
    et = item.get("end_type") or ""
    em = item.get("end_month") or ""
    ey = item.get("end_year") or ""
    if not sy or sy == "None":
        return ""
    start = f"{sm} {sy}" if sm and sm != "None" else sy
    if et == "Present":
        return f"{start} - Present"
    if et == "Specific Month" and ey and ey != "None":
        end = f"{em} {ey}" if em and em != "None" else ey
        return f"{start} - {end}"
    return start


def _add_bottom_border(paragraph, color_hex: str, sz: str):
    """Thin rule under a section heading, matching the chosen template."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color_hex)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _section_title(doc, text: str, style: dict):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    label = (text or "").upper() if style["section_upper"] else (text or "")
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(style["section_size"])
    if style["section_color"]:
        run.font.color.rgb = RGBColor.from_string(style["section_color"])
    _add_bottom_border(p, style["border_color"], style["border_sz"])


def _entry_head(doc, bold_text: str, normal_text: str, dates: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    if dates:
        p.paragraph_format.tab_stops.add_tab_stop(Mm(USABLE_WIDTH_MM), WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(bold_text or "")
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    if dates:
        p.add_run("\t" + dates)


def _bullets(doc, items):
    for b in items or []:
        if b and str(b).strip():
            bp = doc.add_paragraph(str(b).strip(), style="List Bullet")
            bp.paragraph_format.space_after = Pt(0)


def _strip_proto(url: str) -> str:
    return (url or "").strip().replace("https://", "").replace("http://", "").rstrip("/")


def build_docx(resume: dict) -> bytes:
    style = TEMPLATE_STYLES.get(resume.get("template") or "original", TEMPLATE_STYLES["original"])
    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = sec.right_margin = Mm(15)
    sec.top_margin = sec.bottom_margin = Mm(12)

    normal = doc.styles["Normal"]
    normal.font.name = style["font"]
    normal.font.size = Pt(style["body_size"])
    normal.paragraph_format.line_spacing = style["line_spacing"]

    # Header — name, title, contact (aligned per template).
    name_p = doc.add_paragraph()
    name_p.alignment = style["header_align"]
    name_p.paragraph_format.space_after = Pt(0)
    nr = name_p.add_run(resume.get("name") or "")
    nr.bold = True
    nr.font.size = Pt(style["name_size"])

    if resume.get("title"):
        tp = doc.add_paragraph()
        tp.alignment = style["header_align"]
        tp.paragraph_format.space_after = Pt(0)
        tr = tp.add_run(resume["title"])
        tr.font.size = Pt(style["title_size"])
        if style["title_color"]:
            tr.font.color.rgb = RGBColor.from_string(style["title_color"])
            tr.bold = True

    if resume.get("contact_info"):
        cp = doc.add_paragraph()
        cp.alignment = style["header_align"]
        cp.add_run(resume["contact_info"]).font.size = Pt(style["contact_size"])

    for section in resume.get("sections") or []:
        stype = section.get("type")
        title = section.get("title") or ""

        if stype == "paragraph":
            if title:
                _section_title(doc, title, style)
            if section.get("content"):
                doc.add_paragraph(section["content"])

        elif stype == "bullet_points":
            if title:
                _section_title(doc, title, style)
            _bullets(doc, section.get("items"))

        elif stype == "experience":
            if title:
                _section_title(doc, title, style)
            for item in section.get("items") or []:
                company = item.get("company") or ""
                _entry_head(doc, item.get("position") or "", f", {company}" if company else "", _format_date_range(item))
                _bullets(doc, item.get("bullet_points"))

        elif stype == "education":
            if title:
                _section_title(doc, title, style)
            for item in section.get("items") or []:
                inst = item.get("institution") or ""
                _entry_head(doc, item.get("degree") or "", f", {inst}" if inst else "", _format_date_range(item))
                if item.get("details"):
                    doc.add_paragraph(item["details"]).paragraph_format.space_after = Pt(0)

        elif stype == "project":
            if title:
                _section_title(doc, title, style)
            for item in section.get("items") or []:
                _entry_head(doc, item.get("name") or "", "", _format_date_range(item))
                links = []
                if item.get("github"):
                    links.append(_strip_proto(item["github"]))
                if item.get("link"):
                    links.append(_strip_proto(item["link"]))
                if item.get("tech") or links:
                    meta = doc.add_paragraph()
                    meta.paragraph_format.space_after = Pt(0)
                    if item.get("tech"):
                        mr = meta.add_run(f"Stack: {item['tech']}")
                        mr.italic = True
                        mr.font.size = Pt(9.5)
                    if links:
                        if item.get("tech"):
                            meta.add_run("   ")
                        meta.add_run("  ·  ".join(links)).font.size = Pt(9.5)
                _bullets(doc, item.get("bullet_points"))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.post("/generate-docx")
async def generate_docx(resume: dict):
    try:
        data = build_docx(resume)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {e}")

    filename = (resume.get("name") or "resume").replace(" ", "_") + "_Resume.docx"
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
